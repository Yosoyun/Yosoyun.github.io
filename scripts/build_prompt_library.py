from __future__ import annotations

import html
import json
import math
import textwrap
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
DOCS = ROOT / "docs"
DATA = ROOT / "data"

PROJECT_TITLE = "Indrajeet Yadav Math Faculty Prompt Library"
PROJECT_SUBTITLE = (
    "120 copy-paste prompts for mathematics teachers: handwritten solution pages, "
    "multi-method explanations, assignments, DPPs, question papers, rubrics, and revision tools."
)
AUTHOR = "Indrajeet Yadav"


STYLE_PACKS = [
    {
        "name": "Vintage Parchment Masterclass",
        "surface": "warm antique parchment, faint ruled lines, red margin, torn edges",
        "ink": "royal-blue derivation, violet headings, green observations, maroon emphasis",
        "details": "restrained botanical flourishes, tiny stains, museum-quality paper realism",
    },
    {
        "name": "Dark Academia Mathematical Manuscript",
        "surface": "deep ivory handmade paper, sepia wash, ledger ruling, worn book-page edges",
        "ink": "midnight-blue ink, oxblood titles, forest-green marginal notes",
        "details": "classical scrollwork, laurel sprigs, archival annotation marks",
    },
    {
        "name": "Indian Heritage Gurukul Journal",
        "surface": "khadi-style handmade sheet, saffron-beige tone, muted red margin",
        "ink": "deep indigo derivation, aubergine headings, leaf-green hints",
        "details": "subtle lotus-line motifs, mango-leaf curves, dignified hand-drawn separators",
    },
    {
        "name": "Japanese Washi Zen Mathematics",
        "surface": "cream washi paper, visible fibres, feathered edges, soft folded texture",
        "ink": "indigo brush-pen handwriting, plum headings, moss-green notes",
        "details": "sakura petals, bamboo-leaf sketch, sparse red seal-like accent",
    },
    {
        "name": "Renaissance Inventor Codex",
        "surface": "aged cream folio, deckled perimeter, subtle watermark, pale ruling",
        "ink": "cobalt-blue derivation, violet section headings, olive insight notes",
        "details": "compass arcs, proportion sketches, gear-like geometry doodles",
    },
    {
        "name": "Celestial Observatory Notebook",
        "surface": "midnight-blue astronomy paper, silver ruling, soft star-map texture",
        "ink": "white-gold chalk ink, cyan construction lines, violet headings",
        "details": "constellation dots, orbit arcs, precise but quiet astronomical marginalia",
    },
    {
        "name": "Botanical Watercolour Mathematics Journal",
        "surface": "heavy cotton paper, pale wash, delicate deckled edge",
        "ink": "navy main work, berry headings, sage-green hints",
        "details": "tiny watercolor leaves and flowers placed only in empty margins",
    },
    {
        "name": "Architectural Blueprint Study Sheet",
        "surface": "blueprint paper, grid lines, faint measurement marks",
        "ink": "white technical pencil, cyan highlights, amber answer accents",
        "details": "straightedge diagrams, dimension ticks, clean drafting hierarchy",
    },
    {
        "name": "Midnight Black-and-Gold Atelier",
        "surface": "matte black art paper, soft grain, subtle edge light",
        "ink": "ivory ink, muted gold headings, teal hints, rose emphasis",
        "details": "luxury editorial spacing, no glitter, no theatrical darkness",
    },
    {
        "name": "Master Chalkboard Classroom",
        "surface": "deep slate chalkboard, erased ghost marks, faint ruling guides",
        "ink": "soft white and pale-blue chalk, violet headings, green hints",
        "details": "chalk dust, pressure variation, spacious classroom sequencing",
    },
    {
        "name": "Graph Paper Precision Lab",
        "surface": "crisp engineering graph paper with warm daylight shadows",
        "ink": "blue gel pen, red corrections, green side notes",
        "details": "coordinate grids, neat brackets, technical but human handwriting",
    },
    {
        "name": "Premium Smartboard Class Notes",
        "surface": "clean digital smartboard canvas with subtle grid and soft glow",
        "ink": "white stylus writing, cyan arrows, yellow highlights",
        "details": "screen-ready spacing, readable from mobile, no slide-deck clutter",
    },
]


@dataclass
class Category:
    name: str
    short: str
    intent: str
    outputs: str
    scenarios: list[str]
    checks: list[str]


CATEGORIES = [
    Category(
        "Handwritten Solution Image Collections",
        "handwritten",
        "turn one uploaded or pasted problem into premium handwritten solution pages",
        "separate portrait pages, one method per page, same final answer, teacher signature",
        [
            "Five separate handwritten methods for one algebra problem",
            "Five-page calculus limit masterclass",
            "Five trigonometry identity solution pages",
            "Five coordinate geometry approaches with mini-sketches",
            "Five probability solution pages using complementary viewpoints",
            "Five number theory methods for an olympiad-style problem",
            "Five vector or 3D geometry solution pages",
            "Five methods for a quadratic equation or inequality",
            "Five integration approaches with verification",
            "Five-page board-exam solution set from a photographed question",
        ],
        [
            "Each page must contain the full question at the top.",
            "Never combine pages into a collage or grid.",
            "All five pages must reach the identical verified final answer.",
            "If the interface generates one image at a time, generate only the current page.",
        ],
    ),
    Category(
        "Multi-Method Explanations",
        "methods",
        "explain the same question from genuinely different mathematical viewpoints",
        "method names, step-by-step solutions, comparison table, final answer, common mistakes",
        [
            "Three methods for a school-level algebra question",
            "Four methods for a definite integral",
            "Five routes for a geometry proof",
            "Two fast exam methods plus one conceptual method",
            "Graphical, algebraic, and numerical solution comparison",
            "Direct method, substitution method, and identity method",
            "Shortcut method with full proof of why it works",
            "Reverse verification method for MCQ questions",
            "Generalization after solving a specific example",
            "Student-friendly explanation for a hard-looking simple problem",
        ],
        [
            "Different methods must not be cosmetic rearrangements.",
            "Name the key idea before starting each method.",
            "Show where each method is fastest and where it is risky.",
            "End with a final comparison for exam use.",
        ],
    ),
    Category(
        "Assignment Builders",
        "assignment",
        "create ready-to-print classwork, homework, and mixed-practice assignments",
        "instructions, marks, difficulty ladder, answer key, teaching notes, signature footer",
        [
            "Chapter-wise homework assignment with graded difficulty",
            "Mixed-topic revision assignment before a unit test",
            "Concept-building worksheet with examples first",
            "Application-based assignment using real contexts",
            "Error-spotting assignment from common student mistakes",
            "Board-style long-answer assignment",
            "MCQ plus reasoning assignment",
            "Pair-work classroom assignment",
            "Challenge assignment for advanced learners",
            "Remedial assignment for foundational gaps",
        ],
        [
            "Include estimated time and marks.",
            "Avoid duplicate question patterns unless repetition is intentional.",
            "Provide a separate answer key and marking hints.",
            "Balance routine, conceptual, and challenge items.",
        ],
    ),
    Category(
        "DPP and Daily Practice",
        "dpp",
        "design daily practice problems for retention, speed, and concept repair",
        "10-30 question DPP, warm-up, core drill, challenge, homework, answers",
        [
            "15-minute DPP for one subtopic",
            "30-question mixed DPP for weekly revision",
            "Speed drill with answer-only checking",
            "Conceptual DPP with explanation prompts",
            "DPP with three difficulty bands",
            "DPP for weak students after a diagnostic test",
            "DPP for JEE/NEET/foundation-style practice",
            "DPP with spaced repetition from previous chapters",
            "DPP using only NCERT-style language",
            "DPP with one bonus thinking question",
        ],
        [
            "Mark each question with difficulty and target skill.",
            "Keep the DPP finishable in the requested time.",
            "Put answers separately so students do not see them early.",
            "Include two teacher notes on expected mistakes.",
        ],
    ),
    Category(
        "Question Papers and Tests",
        "papers",
        "build complete tests and question papers with clean structure and marking scheme",
        "paper blueprint, sections, marks, time, answer key, rubric, moderation checks",
        [
            "40-mark unit test with balanced sections",
            "80-mark annual exam paper",
            "MCQ-only rapid assessment",
            "Case-study based question paper",
            "Board-style paper with internal choices",
            "Competitive exam practice paper",
            "Chapter diagnostic test",
            "Two-version test with similar difficulty",
            "Open-book conceptual assessment",
            "Exit-ticket mini test for one lesson",
        ],
        [
            "Create a blueprint before the questions.",
            "Ensure marks add up exactly to the requested total.",
            "Avoid ambiguity in instructions and answer choices.",
            "Add a moderation checklist at the end.",
        ],
    ),
    Category(
        "Competitive and Olympiad Training",
        "competitive",
        "prepare higher-order problems, shortcuts, and proof-backed exam strategies",
        "problem set, solution key, traps, alternate methods, time strategy",
        [
            "Olympiad-style inequality practice",
            "JEE-style calculus mixed set",
            "Number theory training ladder",
            "Combinatorics counting strategy set",
            "Coordinate geometry speed practice",
            "Advanced trigonometry transformation set",
            "Functional equation starter pack",
            "Proof without calculus challenge sheet",
            "MCQ elimination strategy set",
            "Hard problem discussion notes for faculty",
        ],
        [
            "Every shortcut must include justification.",
            "Separate elegant solution from exam-speed solution.",
            "Mention prerequisite concepts.",
            "Flag traps that produce attractive wrong answers.",
        ],
    ),
    Category(
        "Board Exam and School Assessment",
        "board",
        "create syllabus-aligned school assessment material with marks and step marking",
        "chapter mapping, expected answer format, marking scheme, common deductions",
        [
            "CBSE-style chapter test",
            "ICSE-style structured answer paper",
            "State board long-answer practice",
            "Step-marking guide for algebra",
            "Assertion-reason practice set",
            "Very short answer bank",
            "Long answer proof question bank",
            "Previous-year pattern remix",
            "Competency-based question set",
            "Revision paper with internal choices",
        ],
        [
            "Use clear school-level language.",
            "Give marks for each major step.",
            "Include answer format students should write.",
            "Avoid changing syllabus boundaries unless asked.",
        ],
    ),
    Category(
        "Remedial Teaching and Doubt Repair",
        "remedial",
        "diagnose student mistakes and rebuild understanding step by step",
        "diagnosis, mini-lesson, practice ladder, teacher script, exit check",
        [
            "Fix misconception in linear equations",
            "Repair sign errors in algebra",
            "Re-teach fraction operations for equations",
            "Doubt-solving script for trigonometric ratios",
            "Error analysis for calculus differentiation",
            "Concept bridge from arithmetic to algebra",
            "Geometry proof confidence builder",
            "Remedial plan after a low test score",
            "Slow learner worksheet with scaffolding",
            "Parent-friendly explanation of a math gap",
        ],
        [
            "Do not shame the student; name the error pattern neutrally.",
            "Start from what the student already understands.",
            "Use tiny steps before full problems.",
            "End with an exit question that reveals whether the misconception is fixed.",
        ],
    ),
    Category(
        "Proof, Concept, and Theory",
        "proof",
        "turn formulas and theorems into clear explanations, proofs, and teachable stories",
        "intuition, formal proof, examples, non-examples, memory hook, final summary",
        [
            "Prove a standard theorem with intuition",
            "Explain why a formula works",
            "Create examples and non-examples",
            "Turn a proof into a classroom dialogue",
            "Concept map for a chapter",
            "Derive formula before using it",
            "Compare two related theorems",
            "Explain hidden assumptions in a result",
            "Create a proof worksheet",
            "Make a one-page theory revision note",
        ],
        [
            "Separate intuition from formal proof.",
            "State all assumptions and domain restrictions.",
            "Include a quick check example.",
            "Avoid hand-waving at the exact step students usually miss.",
        ],
    ),
    Category(
        "Graphs, Diagrams, and Visual Maths",
        "visual",
        "generate graphing instructions, diagrams, and visual explanations that support solutions",
        "diagram specification, labels, step sequence, teacher notes, visual checks",
        [
            "Graph a quadratic and explain roots",
            "Draw a geometry construction with labels",
            "Visualize transformation of functions",
            "Create a coordinate geometry diagram prompt",
            "Explain area under curve visually",
            "Make a vector diagram",
            "Create a number line explanation",
            "Design a probability tree",
            "Build a trigonometry unit-circle visual",
            "Prepare a board diagram plan for class",
        ],
        [
            "Every diagram must have labels and scale notes where needed.",
            "Do not add visual objects that change the mathematics.",
            "Use visuals to explain, not decorate.",
            "Mention what students should notice first.",
        ],
    ),
    Category(
        "Answer Keys, Rubrics, and Checking",
        "rubric",
        "produce reliable answer keys, marking schemes, and proof-checking guides",
        "final answers, step marking, alternate answers, common errors, audit checklist",
        [
            "Detailed answer key for an assignment",
            "Step-marking rubric for a question paper",
            "Fast checking sheet for faculty",
            "Common wrong answers and why they happen",
            "Moderation review for a test paper",
            "Solution verification by two methods",
            "Partial-credit guide for proofs",
            "MCQ key with distractor analysis",
            "Rubric for project-based math work",
            "Student self-check answer sheet",
        ],
        [
            "Verify answers before writing the key.",
            "Give partial marks only for mathematically meaningful steps.",
            "Include alternate valid methods.",
            "Flag ambiguous questions that need revision.",
        ],
    ),
    Category(
        "Revision, Notes, and Faculty Workflow",
        "workflow",
        "help teachers prepare lessons, revision material, and classroom communication",
        "lesson flow, board plan, student handout, revision sheet, follow-up task",
        [
            "One-period lesson plan",
            "Board-work plan for a derivation",
            "Revision sheet before exam",
            "Formula sheet with usage notes",
            "Flashcards for a chapter",
            "Mind map from a chapter list",
            "Parent update for mathematics progress",
            "Student feedback comments from marks",
            "Homework review plan",
            "Faculty meeting note for improving math results",
        ],
        [
            "Make the output immediately usable by a teacher.",
            "Keep classroom timing realistic.",
            "Separate teacher notes from student-facing text.",
            "End with the next action for the class.",
        ],
    ),
]


def slugify(text: str) -> str:
    allowed = []
    for ch in text.lower():
        if ch.isalnum():
            allowed.append(ch)
        elif ch in {" ", "-", "_"}:
            allowed.append("-")
    slug = "".join(allowed)
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug.strip("-")


def build_prompt(idx: int, category: Category, scenario: str, style: dict[str, str]) -> dict[str, str]:
    variables = {
        "GRADE": "Class/grade or exam level",
        "TOPIC": "chapter, subtopic, or skill",
        "SOURCE": "uploaded image, typed question, textbook exercise, or syllabus excerpt",
        "COUNT": "number of questions/pages/methods required",
        "TIME": "available class or practice time",
        "LANGUAGE": "English, Hindi, bilingual, or school-preferred language",
    }
    prompt = f"""ROLE
Act as a senior mathematics faculty member, proof-checker, exam designer, and clear classroom communicator.

USE CASE
{scenario}.

MY INPUTS
- Grade/exam level: [GRADE]
- Topic/chapter: [TOPIC]
- Source material or question: [PASTE_TEXT_OR_UPLOAD_IMAGE]
- Required count/length: [COUNT]
- Time available: [TIME]
- Student level: [WEAK / AVERAGE / ADVANCED / MIXED]
- Language style: [LANGUAGE]

TASK
Create a ready-to-use teacher output for: {category.intent}. The result must be practical for a non-technical mathematics teacher who wants to copy, print, explain, or share the material immediately. Use creator-level clarity: strong hierarchy, precise wording, no filler, and no imitation of any specific living creator.

OUTPUT CONTRACT
1. Start with a short title and the exact target class/topic.
2. State any assumptions you made. If the source question is an image, transcribe it first and solve from that transcription.
3. Produce the requested material: {category.outputs}.
4. Put teacher-facing notes under a separate heading so they are not mixed with student-facing content.
5. Include answer key, final result, or verification wherever relevant.
6. Add a footer line: "Created with the {PROJECT_TITLE} - {AUTHOR}".

QUALITY RULES
- Check the mathematics before finalizing.
- Keep notation consistent from start to finish.
- Use step-by-step explanations without skipping the fragile step students usually miss.
- If there are multiple methods, make them genuinely different.
- Do not invent data, diagrams, answer choices, syllabus rules, or theorem names.
- If the task is too large for one response, finish Part 1 cleanly and ask me to continue with Part 2.

STYLE DIRECTION
Use the visual and instructional mood: {style['name']}. Surface idea: {style['surface']}. Writing/ink idea: {style['ink']}. Detail discipline: {style['details']}. For text-only outputs, translate this into clean headings, spacing, and marginal notes instead of decoration.

FINAL SELF-CHECK
Before giving the answer, verify:
- The output matches [GRADE], [TOPIC], [COUNT], and [TIME].
- The answer key or final answer is mathematically consistent.
- Instructions are clear enough for a teacher to use without technical knowledge.
- The {AUTHOR} project credit appears exactly once at the end.
"""
    return {
        "id": f"P{idx:03d}",
        "title": scenario,
        "category": category.name,
        "category_short": category.short,
        "style": style["name"],
        "use_case": category.intent,
        "variables": ", ".join(f"[{key}]" for key in variables),
        "prompt": textwrap.dedent(prompt).strip(),
    }


def generate_prompts() -> list[dict[str, str]]:
    prompts: list[dict[str, str]] = []
    idx = 1
    for cat_index, category in enumerate(CATEGORIES):
        for scenario_index, scenario in enumerate(category.scenarios):
            style = STYLE_PACKS[(cat_index * 3 + scenario_index) % len(STYLE_PACKS)]
            prompts.append(build_prompt(idx, category, scenario, style))
            idx += 1
    return prompts


def make_preview_image(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1400, 900
    img = Image.new("RGB", (w, h), "#f6f1e7")
    draw = ImageDraw.Draw(img)
    for y in range(h):
        tint = int(10 * math.sin(y / 70))
        draw.line((0, y, w, y), fill=(246 + tint // 3, 241 + tint // 5, 231 + tint // 6))
    for x in range(0, w, 42):
        draw.line((x, 0, x, h), fill="#ebe1d2", width=1)
    for y in range(0, h, 42):
        draw.line((0, y, w, y), fill="#ebe1d2", width=1)
    fonts = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Georgia.ttf",
    ]
    font_path = next((p for p in fonts if Path(p).exists()), None)
    title_font = ImageFont.truetype(font_path, 54) if font_path else ImageFont.load_default()
    sub_font = ImageFont.truetype(font_path, 28) if font_path else ImageFont.load_default()
    mono_font = ImageFont.truetype(font_path, 24) if font_path else ImageFont.load_default()

    draw.rounded_rectangle((70, 70, 1330, 830), radius=34, fill="#fffaf1", outline="#d3b77c", width=4)
    draw.rectangle((120, 120, 124, 790), fill="#b23a48")
    draw.text((155, 125), "Math Faculty Prompt Library", fill="#13294b", font=title_font)
    draw.text((158, 198), "120 copy-paste prompts by Indrajeet Yadav", fill="#40635b", font=sub_font)
    draw.line((155, 255, 1200, 255), fill="#d8b55b", width=4)
    sample_lines = [
        "ROLE: Act as a senior mathematics educator...",
        "TASK: Build five different solution methods...",
        "OUTPUT: question, steps, checks, answer key...",
        "SIGNATURE: Indrajeet Yadav",
    ]
    y = 315
    colors = ["#174a7c", "#7a2f6f", "#2f6b4f", "#8d3b36"]
    for line, color in zip(sample_lines, colors):
        draw.text((170, y), line, fill=color, font=mono_font)
        y += 62

    for i, (x, y0, c) in enumerate([(840, 340, "#efe3c8"), (920, 400, "#dbe7e1"), (990, 310, "#eadbe6")]):
        draw.rounded_rectangle((x, y0, x + 210, y0 + 260), radius=20, fill=c, outline="#bfa46d", width=2)
        draw.line((x + 22, y0 + 52, x + 185, y0 + 52), fill="#13294b", width=3)
        for k in range(5):
            draw.line((x + 24, y0 + 90 + k * 30, x + 178, y0 + 90 + k * 30), fill="#174a7c", width=2)
        draw.text((x + 55, y0 + 205), f"P{i+1:02d}", fill="#8d3b36", font=sub_font)

    draw.text((920, 760), "Indrajeet Yadav", fill="#13294b", font=sub_font)
    draw.line((918, 800, 1175, 800), fill="#13294b", width=2)
    img.save(path, quality=94)


def render_markdown(prompts: list[dict[str, str]]) -> str:
    lines = [
        f"# {PROJECT_TITLE}",
        "",
        PROJECT_SUBTITLE,
        "",
        f"Created and authored by {AUTHOR}.",
        "",
        "## How teachers should use this pack",
        "",
        "1. Open ChatGPT.",
        "2. Choose one prompt from this library.",
        "3. Replace the square-bracket fields such as [GRADE], [TOPIC], and [COUNT].",
        "4. Upload a mathematics problem image only when the prompt asks for it.",
        "5. Paste the prompt and read the output carefully.",
        "6. Ask ChatGPT to verify the answer or simplify the language if needed.",
        "7. For image prompts, generate one page at a time if the interface cannot return five files together.",
        "",
        "## Master input block",
        "",
        "```text",
        "Grade/exam level: [GRADE]",
        "Topic/chapter: [TOPIC]",
        "Source material/question: [PASTE_TEXT_OR_UPLOAD_IMAGE]",
        "Required count/length: [COUNT]",
        "Time available: [TIME]",
        "Student level: [WEAK / AVERAGE / ADVANCED / MIXED]",
        "Language style: [LANGUAGE]",
        "```",
        "",
        "## Prompt Library",
        "",
    ]
    current = None
    for prompt in prompts:
        if prompt["category"] != current:
            current = prompt["category"]
            lines.extend(["", f"## {current}", ""])
        lines.extend(
            [
                f"### {prompt['id']} - {prompt['title']}",
                "",
                f"Use case: {prompt['use_case']}",
                "",
                f"Style pack: {prompt['style']}",
                "",
                "```text",
                prompt["prompt"],
                "```",
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n"


def render_html(prompts: list[dict[str, str]]) -> str:
    categories = sorted({p["category"] for p in prompts})
    data_json = json.dumps(prompts, ensure_ascii=False)
    style_json = json.dumps(STYLE_PACKS, ensure_ascii=False)
    category_buttons = "\n".join(
        f'<button class="chip" data-filter="{html.escape(cat)}">{html.escape(cat)}</button>' for cat in categories
    )
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(PROJECT_TITLE)}</title>
  <meta name="description" content="{html.escape(PROJECT_SUBTITLE)}">
  <link rel="preload" as="image" href="assets/prompt-library-preview.png">
  <style>
    :root {{
      --ink: #14213d;
      --muted: #53606f;
      --paper: #fffaf1;
      --wash: #f5efe3;
      --line: #d8c7a3;
      --gold: #b88a24;
      --green: #2f6b4f;
      --maroon: #9b3345;
      --blue: #17518a;
      --plum: #6d3b73;
      --shadow: 0 16px 50px rgba(20, 33, 61, .13);
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      color: var(--ink);
      background:
        linear-gradient(90deg, rgba(184,138,36,.08) 1px, transparent 1px),
        linear-gradient(0deg, rgba(20,33,61,.05) 1px, transparent 1px),
        #fbf8f1;
      background-size: 46px 46px;
      letter-spacing: 0;
    }}
    a {{ color: var(--blue); }}
    .topbar {{
      position: sticky;
      top: 0;
      z-index: 10;
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 16px;
      padding: 12px clamp(18px, 4vw, 56px);
      background: rgba(255, 250, 241, .94);
      border-bottom: 1px solid rgba(216, 199, 163, .78);
      backdrop-filter: blur(12px);
    }}
    .brand {{ display: flex; align-items: center; gap: 12px; font-weight: 800; }}
    .brandmark {{
      width: 36px;
      height: 36px;
      display: grid;
      place-items: center;
      border-radius: 8px;
      background: var(--ink);
      color: #fffaf1;
      font-weight: 900;
    }}
    .navlinks {{ display: flex; gap: 12px; flex-wrap: wrap; justify-content: flex-end; }}
    .navlinks a {{ font-size: 14px; font-weight: 700; text-decoration: none; }}
    .hero {{
      min-height: calc(100vh - 64px);
      display: grid;
      grid-template-columns: minmax(0, 1.02fr) minmax(300px, .98fr);
      align-items: center;
      gap: clamp(28px, 5vw, 70px);
      padding: clamp(28px, 5vw, 72px) clamp(18px, 4vw, 56px) 40px;
    }}
    .eyebrow {{
      display: inline-flex;
      align-items: center;
      gap: 8px;
      color: var(--maroon);
      font-weight: 800;
      text-transform: uppercase;
      font-size: 12px;
      margin-bottom: 16px;
    }}
    h1 {{
      margin: 0;
      max-width: 780px;
      font-size: clamp(38px, 6vw, 82px);
      line-height: .98;
      letter-spacing: 0;
    }}
    .lead {{
      max-width: 740px;
      color: var(--muted);
      font-size: clamp(18px, 2.3vw, 24px);
      line-height: 1.45;
      margin: 20px 0 24px;
    }}
    .hero-actions {{ display: flex; gap: 12px; flex-wrap: wrap; }}
    .button {{
      border: 1px solid var(--ink);
      background: var(--ink);
      color: white;
      padding: 12px 16px;
      border-radius: 8px;
      font-weight: 800;
      text-decoration: none;
      display: inline-flex;
      align-items: center;
      gap: 8px;
      min-height: 44px;
    }}
    .button.secondary {{ background: transparent; color: var(--ink); border-color: var(--line); }}
    .hero-visual img {{
      width: 100%;
      max-height: 70vh;
      object-fit: contain;
      border-radius: 8px;
      box-shadow: var(--shadow);
      border: 1px solid var(--line);
      background: #fff;
    }}
    .stats {{
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 10px;
      margin-top: 22px;
      max-width: 680px;
    }}
    .stat {{
      border: 1px solid var(--line);
      background: rgba(255,250,241,.72);
      padding: 14px;
      border-radius: 8px;
    }}
    .stat strong {{ display: block; font-size: 24px; color: var(--blue); }}
    section {{
      padding: 54px clamp(18px, 4vw, 56px);
      border-top: 1px solid rgba(216, 199, 163, .7);
      background: rgba(255, 250, 241, .66);
    }}
    section.alt {{ background: rgba(238, 245, 241, .76); }}
    .section-head {{
      display: flex;
      justify-content: space-between;
      align-items: end;
      gap: 24px;
      margin-bottom: 24px;
    }}
    h2 {{ margin: 0; font-size: clamp(26px, 3vw, 42px); letter-spacing: 0; }}
    .section-note {{ max-width: 780px; color: var(--muted); line-height: 1.55; }}
    .steps {{
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      gap: 14px;
    }}
    .step, .download, .style-card, .prompt-card {{
      background: var(--paper);
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: 0 8px 28px rgba(20,33,61,.07);
    }}
    .step {{ padding: 18px; min-height: 150px; }}
    .step b {{ color: var(--maroon); }}
    .step p {{ color: var(--muted); line-height: 1.5; margin: 10px 0 0; }}
    .download-grid, .style-grid {{
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      gap: 14px;
    }}
    .download, .style-card {{ padding: 18px; }}
    .download h3, .style-card h3 {{ margin: 0 0 8px; font-size: 18px; }}
    .download p, .style-card p {{ margin: 0; color: var(--muted); line-height: 1.45; }}
    .controls {{
      display: grid;
      grid-template-columns: minmax(220px, 1fr) auto;
      gap: 12px;
      align-items: center;
      margin-bottom: 16px;
    }}
    .search {{
      width: 100%;
      min-height: 48px;
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 0 14px;
      font-size: 16px;
      background: white;
    }}
    .chips {{ display: flex; flex-wrap: wrap; gap: 8px; margin-bottom: 18px; }}
    .chip {{
      min-height: 38px;
      border: 1px solid var(--line);
      background: #fff;
      border-radius: 8px;
      padding: 8px 11px;
      cursor: pointer;
      color: var(--ink);
      font-weight: 700;
      font-size: 13px;
    }}
    .chip.active {{ background: var(--green); color: white; border-color: var(--green); }}
    .library-count {{ color: var(--muted); font-weight: 800; white-space: nowrap; }}
    .prompt-grid {{
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 16px;
      align-items: start;
    }}
    .prompt-card {{
      padding: 18px;
      display: flex;
      flex-direction: column;
      gap: 12px;
      min-height: 430px;
    }}
    .prompt-meta {{
      display: flex;
      justify-content: space-between;
      gap: 8px;
      color: var(--muted);
      font-size: 12px;
      font-weight: 800;
      text-transform: uppercase;
    }}
    .prompt-card h3 {{ margin: 0; font-size: 20px; line-height: 1.2; }}
    .prompt-card .use {{ color: var(--muted); line-height: 1.45; margin: 0; }}
    pre {{
      white-space: pre-wrap;
      overflow: auto;
      max-height: 270px;
      margin: 0;
      padding: 14px;
      background: #fbf7ee;
      border: 1px solid #ead9ba;
      border-radius: 8px;
      color: #16324f;
      font-size: 12px;
      line-height: 1.45;
      font-family: "SFMono-Regular", Consolas, "Liberation Mono", monospace;
    }}
    .card-actions {{ display: flex; justify-content: space-between; gap: 10px; margin-top: auto; }}
    .copy-btn, .expand-btn {{
      min-height: 40px;
      border-radius: 8px;
      border: 1px solid var(--line);
      padding: 8px 12px;
      background: white;
      color: var(--ink);
      font-weight: 800;
      cursor: pointer;
    }}
    .copy-btn {{ background: var(--blue); color: white; border-color: var(--blue); }}
    .prompt-card.expanded pre {{ max-height: none; }}
    .master-block {{
      display: grid;
      grid-template-columns: minmax(0, 1fr) minmax(280px, 420px);
      gap: 22px;
      align-items: start;
    }}
    .signature {{
      font-family: Georgia, serif;
      font-style: italic;
      font-size: 28px;
      color: var(--maroon);
      border-bottom: 2px solid var(--maroon);
      display: inline-block;
      padding-bottom: 4px;
    }}
    footer {{
      padding: 34px clamp(18px, 4vw, 56px);
      background: var(--ink);
      color: #fffaf1;
    }}
    footer p {{ margin: 6px 0; color: #e8dcc8; }}
    @media (max-width: 1040px) {{
      .hero, .master-block {{ grid-template-columns: 1fr; }}
      .steps, .download-grid, .style-grid, .prompt-grid {{ grid-template-columns: repeat(2, minmax(0, 1fr)); }}
    }}
    @media (max-width: 680px) {{
      .topbar {{ position: static; align-items: flex-start; flex-direction: column; }}
      .stats, .steps, .download-grid, .style-grid, .prompt-grid, .controls {{ grid-template-columns: 1fr; }}
      .hero {{ min-height: auto; padding-top: 30px; }}
      h1 {{ font-size: 40px; }}
      pre {{ max-height: 230px; }}
    }}
  </style>
</head>
<body>
  <header class="topbar">
    <div class="brand"><span class="brandmark">IY</span><span>{html.escape(PROJECT_TITLE)}</span></div>
    <nav class="navlinks" aria-label="Page navigation">
      <a href="#start">Start</a>
      <a href="#downloads">Downloads</a>
      <a href="#library">Prompts</a>
      <a href="#styles">Styles</a>
    </nav>
  </header>

  <main>
    <section class="hero" id="start">
      <div>
        <div class="eyebrow">Copy-paste classroom prompt system</div>
        <h1>{html.escape(PROJECT_TITLE)}</h1>
        <p class="lead">{html.escape(PROJECT_SUBTITLE)} Built for non-technical faculty who want better ChatGPT results without learning APIs, tools, or complicated prompt engineering.</p>
        <div class="hero-actions">
          <a class="button" href="#library">Browse 120 prompts</a>
          <a class="button secondary" href="prompt-pack.md">Open Markdown pack</a>
          <a class="button secondary" href="docs/Indrajeet-Yadav-Math-Faculty-Prompt-Library.pdf">Download PDF</a>
        </div>
        <div class="stats">
          <div class="stat"><strong>120</strong><span>ready prompts</span></div>
          <div class="stat"><strong>12</strong><span>teacher workflows</span></div>
          <div class="stat"><strong>5x</strong><span>multi-method solution support</span></div>
        </div>
      </div>
      <div class="hero-visual">
        <img src="assets/prompt-library-preview.png" alt="Preview of handwritten mathematics prompt library pages">
      </div>
    </section>

    <section class="alt">
      <div class="section-head">
        <div>
          <h2>How To Use</h2>
          <p class="section-note">Every prompt is already written. Teachers only replace the square brackets, paste into ChatGPT, and review the result before using it with students.</p>
        </div>
      </div>
      <div class="steps">
        <div class="step"><b>1. Pick</b><p>Choose one card matching your task: solution, assignment, DPP, paper, rubric, or revision.</p></div>
        <div class="step"><b>2. Replace</b><p>Fill [GRADE], [TOPIC], [COUNT], [TIME], and paste or upload the question if needed.</p></div>
        <div class="step"><b>3. Paste</b><p>Send the prompt in ChatGPT. For image pages, generate one page at a time if needed.</p></div>
        <div class="step"><b>4. Verify</b><p>Ask for a quick self-check, confirm the answer key, then print, share, or adapt.</p></div>
      </div>
    </section>

    <section id="downloads">
      <div class="section-head">
        <div>
          <h2>Downloads</h2>
          <p class="section-note">The same library is available as a webpage, Markdown pack, PDF, DOCX, and structured JSON.</p>
        </div>
      </div>
      <div class="download-grid">
        <a class="download" href="prompt-pack.md"><h3>Markdown Prompt Pack</h3><p>Best for direct copying, editing, and GitHub reading.</p></a>
        <a class="download" href="docs/Indrajeet-Yadav-Math-Faculty-Prompt-Library.pdf"><h3>PDF Guide</h3><p>Printable teacher handout with all prompts and usage steps.</p></a>
        <a class="download" href="docs/Indrajeet-Yadav-Math-Faculty-Prompt-Library.docx"><h3>Word DOCX</h3><p>Editable document for faculty training or school customization.</p></a>
        <a class="download" href="data/prompts.json"><h3>Prompt JSON</h3><p>Structured data for future apps, search, or expansion.</p></a>
      </div>
    </section>

    <section class="alt">
      <div class="master-block">
        <div>
          <h2>Master Input Block</h2>
          <p class="section-note">Faculty can fill this once, then reuse it with any prompt. It prevents vague outputs and makes ChatGPT behave like a focused teaching assistant.</p>
          <pre id="masterPrompt">Grade/exam level: [GRADE]
Topic/chapter: [TOPIC]
Source material/question: [PASTE_TEXT_OR_UPLOAD_IMAGE]
Required count/length: [COUNT]
Time available: [TIME]
Student level: [WEAK / AVERAGE / ADVANCED / MIXED]
Language style: [LANGUAGE]</pre>
          <p><button class="copy-btn" data-copy-target="masterPrompt">Copy master block</button></p>
        </div>
        <div>
          <h2>Project Credit</h2>
          <p class="section-note">Each prompt includes a footer credit so teachers remember the project source while keeping their classroom material usable.</p>
          <p class="signature">{html.escape(AUTHOR)}</p>
        </div>
      </div>
    </section>

    <section id="library">
      <div class="section-head">
        <div>
          <h2>Prompt Library</h2>
          <p class="section-note">Search by topic, category, use case, or style. Open a card, copy the prompt, then replace the bracketed fields.</p>
        </div>
        <div class="library-count" id="countLabel">120 prompts</div>
      </div>
      <div class="controls">
        <input class="search" id="searchBox" type="search" placeholder="Search prompts, categories, or styles">
        <button class="chip active" data-filter="all">All prompts</button>
      </div>
      <div class="chips" id="categoryChips">
        {category_buttons}
      </div>
      <div class="prompt-grid" id="promptGrid"></div>
    </section>

    <section class="alt" id="styles">
      <div class="section-head">
        <div>
          <h2>Style Vault</h2>
          <p class="section-note">These style packs are included inside the prompts so faculty can request premium handwritten pages, notes, worksheets, and visual explanations without writing style instructions from scratch.</p>
        </div>
      </div>
      <div class="style-grid" id="styleGrid"></div>
    </section>
  </main>

  <footer>
    <strong>{html.escape(PROJECT_TITLE)}</strong>
    <p>Created and authored by {html.escape(AUTHOR)}. Free for mathematics faculty to use, adapt, and share with credit.</p>
  </footer>

  <script>
    const PROMPTS = {data_json};
    const STYLES = {style_json};
    const grid = document.querySelector("#promptGrid");
    const searchBox = document.querySelector("#searchBox");
    const countLabel = document.querySelector("#countLabel");
    let activeFilter = "all";

    function escapeHtml(value) {{
      return value.replace(/[&<>"']/g, ch => ({{'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}}[ch]));
    }}

    function renderPrompts() {{
      const query = searchBox.value.trim().toLowerCase();
      const filtered = PROMPTS.filter(p => {{
        const haystack = `${{p.id}} ${{p.title}} ${{p.category}} ${{p.style}} ${{p.use_case}} ${{p.prompt}}`.toLowerCase();
        const matchesSearch = !query || haystack.includes(query);
        const matchesFilter = activeFilter === "all" || p.category === activeFilter;
        return matchesSearch && matchesFilter;
      }});
      countLabel.textContent = `${{filtered.length}} prompt${{filtered.length === 1 ? "" : "s"}}`;
      grid.innerHTML = filtered.map(p => `
        <article class="prompt-card" data-id="${{p.id}}">
          <div class="prompt-meta"><span>${{p.id}}</span><span>${{escapeHtml(p.category_short)}}</span></div>
          <h3>${{escapeHtml(p.title)}}</h3>
          <p class="use">${{escapeHtml(p.use_case)}}</p>
          <div class="prompt-meta"><span>${{escapeHtml(p.style)}}</span></div>
          <pre id="prompt-${{p.id}}">${{escapeHtml(p.prompt)}}</pre>
          <div class="card-actions">
            <button class="copy-btn" data-copy-target="prompt-${{p.id}}">Copy prompt</button>
            <button class="expand-btn" data-expand="${{p.id}}">Expand</button>
          </div>
        </article>
      `).join("");
    }}

    function renderStyles() {{
      document.querySelector("#styleGrid").innerHTML = STYLES.map(style => `
        <article class="style-card">
          <h3>${{escapeHtml(style.name)}}</h3>
          <p><b>Surface:</b> ${{escapeHtml(style.surface)}}</p>
          <p><b>Ink:</b> ${{escapeHtml(style.ink)}}</p>
          <p><b>Details:</b> ${{escapeHtml(style.details)}}</p>
        </article>
      `).join("");
    }}

    document.addEventListener("click", async event => {{
      const copyButton = event.target.closest("[data-copy-target]");
      if (copyButton) {{
        const target = document.getElementById(copyButton.dataset.copyTarget);
        await navigator.clipboard.writeText(target.textContent);
        const original = copyButton.textContent;
        copyButton.textContent = "Copied";
        setTimeout(() => copyButton.textContent = original, 1200);
      }}
      const expandButton = event.target.closest("[data-expand]");
      if (expandButton) {{
        const card = expandButton.closest(".prompt-card");
        card.classList.toggle("expanded");
        expandButton.textContent = card.classList.contains("expanded") ? "Collapse" : "Expand";
      }}
      const filterButton = event.target.closest("[data-filter]");
      if (filterButton) {{
        activeFilter = filterButton.dataset.filter;
        document.querySelectorAll("[data-filter]").forEach(btn => btn.classList.toggle("active", btn === filterButton));
        renderPrompts();
      }}
    }});
    searchBox.addEventListener("input", renderPrompts);
    renderPrompts();
    renderStyles();
  </script>
</body>
</html>
"""


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)


def shade_paragraph(paragraph, fill: str = "FBF7EE") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def create_docx(prompts: list[dict[str, str]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"{PROJECT_TITLE} | Created by {AUTHOR}")
    set_run_font(r, size=9, color="53606F")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(PROJECT_TITLE)
    set_run_font(run, size=26, color="14213D", bold=True)
    subtitle = doc.add_paragraph(PROJECT_SUBTITLE)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.line_spacing = 1.25
    set_run_font(subtitle.runs[0], size=12, color="53606F")
    byline = doc.add_paragraph(f"Created and authored by {AUTHOR}.")
    set_run_font(byline.runs[0], size=11, color="9B3345", bold=True)

    doc.add_heading("How teachers should use this pack", level=1)
    for item in [
        "Open ChatGPT and choose one prompt from this library.",
        "Replace square-bracket fields such as [GRADE], [TOPIC], [COUNT], and [TIME].",
        "Upload a mathematics problem image only when the prompt asks for it.",
        "Paste the prompt and review the answer carefully before using it with students.",
        "For image prompts, generate one page at a time if the interface cannot return five files together.",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

    doc.add_heading("Master input block", level=1)
    block = doc.add_paragraph()
    shade_paragraph(block)
    block.paragraph_format.space_before = Pt(4)
    block.paragraph_format.space_after = Pt(10)
    block.paragraph_format.line_spacing = 1.15
    r = block.add_run(
        "Grade/exam level: [GRADE]\n"
        "Topic/chapter: [TOPIC]\n"
        "Source material/question: [PASTE_TEXT_OR_UPLOAD_IMAGE]\n"
        "Required count/length: [COUNT]\n"
        "Time available: [TIME]\n"
        "Student level: [WEAK / AVERAGE / ADVANCED / MIXED]\n"
        "Language style: [LANGUAGE]"
    )
    set_run_font(r, name="Courier New", size=9.5, color="16324F")

    doc.add_heading("Library index", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(1.0), Inches(3.1), Inches(2.4)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "ID", True)
    set_cell_text(hdr[1], "Prompt", True)
    set_cell_text(hdr[2], "Category", True)
    for prompt in prompts:
        cells = table.add_row().cells
        set_cell_text(cells[0], prompt["id"])
        set_cell_text(cells[1], prompt["title"])
        set_cell_text(cells[2], prompt["category"])

    doc.add_page_break()
    current = None
    for prompt in prompts:
        if prompt["category"] != current:
            if current is not None:
                doc.add_page_break()
            current = prompt["category"]
            doc.add_heading(current, level=1)
        doc.add_heading(f"{prompt['id']} - {prompt['title']}", level=2)
        meta = doc.add_paragraph(f"Use case: {prompt['use_case']} | Style pack: {prompt['style']}")
        set_run_font(meta.runs[0], size=9.5, color="53606F", italic=True)
        p = doc.add_paragraph()
        shade_paragraph(p)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(prompt["prompt"])
        set_run_font(r, name="Courier New", size=8.2, color="16324F")

    doc.save(path)


def wrap_text(text: str, width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.splitlines():
        if not raw:
            lines.append("")
            continue
        lines.extend(textwrap.wrap(raw, width=width, replace_whitespace=False, drop_whitespace=False))
    return lines


def create_pdf(prompts: list[dict[str, str]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=letter)
    w, h = letter
    margin = 54
    y = h - margin
    page = 1

    def footer() -> None:
        c.setStrokeColorRGB(0.82, 0.74, 0.58)
        c.line(margin, 36, w - margin, 36)
        c.setFillColorRGB(0.32, 0.38, 0.44)
        c.setFont("Helvetica", 8)
        c.drawString(margin, 22, f"{PROJECT_TITLE} | Created by {AUTHOR}")
        c.drawRightString(w - margin, 22, f"Page {page}")

    def new_page() -> None:
        nonlocal y, page
        footer()
        c.showPage()
        page += 1
        y = h - margin

    def draw_wrapped(text: str, font: str = "Helvetica", size: int = 10, color=(0.08, 0.13, 0.24), leading: int = 13, width_chars: int = 95) -> None:
        nonlocal y
        c.setFont(font, size)
        c.setFillColorRGB(*color)
        for line in wrap_text(text, width_chars):
            if y < 64:
                new_page()
                c.setFont(font, size)
                c.setFillColorRGB(*color)
            c.drawString(margin, y, line)
            y -= leading

    c.setFillColorRGB(0.08, 0.13, 0.24)
    c.setFont("Helvetica-Bold", 24)
    c.drawString(margin, y, PROJECT_TITLE)
    y -= 34
    c.setFont("Helvetica", 11)
    c.setFillColorRGB(0.32, 0.38, 0.44)
    draw_wrapped(PROJECT_SUBTITLE, size=11, leading=15, width_chars=88)
    y -= 8
    c.setFont("Helvetica-Bold", 11)
    c.setFillColorRGB(0.61, 0.20, 0.27)
    c.drawString(margin, y, f"Created and authored by {AUTHOR}.")
    y -= 28
    c.setFillColorRGB(0.18, 0.42, 0.31)
    c.setFont("Helvetica-Bold", 15)
    c.drawString(margin, y, "How teachers should use this pack")
    y -= 22
    for step in [
        "1. Open ChatGPT.",
        "2. Choose one prompt from this library.",
        "3. Replace the square-bracket fields such as [GRADE], [TOPIC], and [COUNT].",
        "4. Upload a mathematics problem image only when the prompt asks for it.",
        "5. Paste the prompt and review the result before using it with students.",
        "6. For image prompts, generate one page at a time if the interface cannot return five files together.",
    ]:
        draw_wrapped(step, size=10, leading=14, width_chars=92)
    y -= 12

    current = None
    for prompt in prompts:
        if prompt["category"] != current:
            current = prompt["category"]
            if y < 180:
                new_page()
            c.setFillColorRGB(0.09, 0.32, 0.54)
            c.setFont("Helvetica-Bold", 16)
            c.drawString(margin, y, current)
            y -= 24
        if y < 160:
            new_page()
        c.setFillColorRGB(0.08, 0.13, 0.24)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(margin, y, f"{prompt['id']} - {prompt['title']}")
        y -= 16
        draw_wrapped(f"Use case: {prompt['use_case']} | Style pack: {prompt['style']}", size=8.5, color=(0.32, 0.38, 0.44), leading=11, width_chars=108)
        y -= 2
        draw_wrapped(prompt["prompt"], font="Courier", size=7.2, color=(0.08, 0.20, 0.31), leading=9, width_chars=110)
        y -= 12

    footer()
    c.save()


def create_readme(prompts: list[dict[str, str]]) -> str:
    return f"""# {PROJECT_TITLE}

{PROJECT_SUBTITLE}

Created and authored by **{AUTHOR}**.

## What is inside

- `index.html` - public copy-button webpage.
- `prompt-pack.md` - complete Markdown prompt pack.
- `docs/Indrajeet-Yadav-Math-Faculty-Prompt-Library.pdf` - printable PDF guide.
- `docs/Indrajeet-Yadav-Math-Faculty-Prompt-Library.docx` - editable Word document.
- `data/prompts.json` - structured source data for all {len(prompts)} prompts.

## How to use

Open `index.html`, choose a prompt, replace the square-bracket fields, and paste it into ChatGPT. The prompts are designed for non-technical mathematics faculty and do not require APIs or external tools.

## Project credit

Please keep the footer credit in shared outputs:

`Created with the {PROJECT_TITLE} - {AUTHOR}`

## GitHub Pages

This repository is static. It can be served directly from the root of a GitHub Pages branch.
"""


def main() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    DOCS.mkdir(parents=True, exist_ok=True)
    DATA.mkdir(parents=True, exist_ok=True)
    prompts = generate_prompts()
    (DATA / "prompts.json").write_text(json.dumps(prompts, indent=2, ensure_ascii=False), encoding="utf-8")
    (ROOT / "prompt-pack.md").write_text(render_markdown(prompts), encoding="utf-8")
    (ROOT / "index.html").write_text(render_html(prompts), encoding="utf-8")
    (ROOT / "README.md").write_text(create_readme(prompts), encoding="utf-8")
    (ROOT / ".nojekyll").write_text("", encoding="utf-8")
    make_preview_image(ASSETS / "prompt-library-preview.png")
    create_docx(prompts, DOCS / "Indrajeet-Yadav-Math-Faculty-Prompt-Library.docx")
    create_pdf(prompts, DOCS / "Indrajeet-Yadav-Math-Faculty-Prompt-Library.pdf")
    print(f"Generated {len(prompts)} prompts in {ROOT}")


if __name__ == "__main__":
    main()
